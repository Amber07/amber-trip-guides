#!/usr/bin/env python3
"""从常见行程文件中抽取纯文本，供 Codex 结构化成 trip-data.json。"""

from __future__ import annotations

import argparse
import csv
import shutil
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def read_csv_like(path: Path, delimiter: str) -> str:
    rows: list[str] = []
    with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as handle:
      reader = csv.reader(handle, delimiter=delimiter)
      for row in reader:
          rows.append(" | ".join(cell.strip() for cell in row))
    return "\n".join(rows)


def safe_name(value: str) -> str:
    stem = "".join(ch if ch.isalnum() or ch in "-_." else "_" for ch in value)
    return stem.strip("._") or "asset"


def extract_zip_media(path: Path, media_dir: Path, prefixes: tuple[str, ...]) -> list[Path]:
    saved: list[Path] = []
    media_dir.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path) as archive:
        for name in archive.namelist():
            if not name.startswith(prefixes) or name.endswith("/"):
                continue
            suffix = Path(name).suffix.lower()
            if suffix not in {".jpg", ".jpeg", ".png", ".webp", ".gif"}:
                continue
            target = media_dir / f"{path.stem}_{safe_name(Path(name).name)}"
            with archive.open(name) as source, target.open("wb") as handle:
                shutil.copyfileobj(source, handle)
            saved.append(target)
    return saved


def extract_pdf_media(path: Path, media_dir: Path) -> list[Path]:
    try:
        import fitz
    except ModuleNotFoundError:
        return []

    saved: list[Path] = []
    media_dir.mkdir(parents=True, exist_ok=True)
    document = fitz.open(path)
    for page_index in range(len(document)):
        for image_index, image in enumerate(document[page_index].get_images(full=True), start=1):
            xref = image[0]
            payload = document.extract_image(xref)
            ext = payload.get("ext", "png")
            target = media_dir / f"{path.stem}_p{page_index + 1:02d}_{image_index:02d}.{ext}"
            target.write_bytes(payload["image"])
            saved.append(target)
    return saved


def read_xlsx(path: Path) -> str:
    try:
        import openpyxl
    except ModuleNotFoundError as exc:
        raise RuntimeError("缺少 openpyxl，无法读取 Excel 文件") from exc

    workbook = openpyxl.load_workbook(path, data_only=True, read_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if cell is None else str(cell).strip() for cell in row]
            if any(values):
                parts.append(" | ".join(values))
    return "\n".join(parts)


def read_docx(path: Path) -> str:
    try:
        import docx
        document = docx.Document(path)
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)
    except ModuleNotFoundError:
        # 无 python-docx 时，直接从 OOXML 解压正文，覆盖普通 docx 文本。
        with zipfile.ZipFile(path) as archive:
            xml = archive.read("word/document.xml")
        root = ElementTree.fromstring(xml)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        paragraphs: list[str] = []
        for para in root.findall(".//w:p", ns):
            text = "".join(node.text or "" for node in para.findall(".//w:t", ns)).strip()
            if text:
                paragraphs.append(text)
        return "\n".join(paragraphs)


def read_pdf(path: Path) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return "\n".join((page.extract_text() or "") for page in pdf.pages)
    except ModuleNotFoundError:
        pass

    try:
        from pypdf import PdfReader
    except ModuleNotFoundError as exc:
        raise RuntimeError("缺少 pdfplumber 或 pypdf，无法读取 PDF 文件") from exc

    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def media_note(saved: list[Path], media_dir: Path, media_path_prefix: str | None) -> str:
    if not saved:
        return ""
    prefix = media_path_prefix if media_path_prefix is not None else media_dir.name
    refs = [str(Path(prefix) / item.name) if prefix else item.name for item in saved]
    return "\n[文档图片]\n" + "\n".join(refs)


def extract(path: Path, media_dir: Path | None = None, media_path_prefix: str | None = None) -> str:
    suffix = path.suffix.lower()
    note = ""
    if media_dir and suffix == ".docx":
        saved = extract_zip_media(path, media_dir, ("word/media/",))
        note = media_note(saved, media_dir, media_path_prefix)
    if media_dir and suffix in {".xlsx", ".xlsm"}:
        saved = extract_zip_media(path, media_dir, ("xl/media/",))
        note = media_note(saved, media_dir, media_path_prefix)
    if media_dir and suffix == ".pdf":
        saved = extract_pdf_media(path, media_dir)
        note = media_note(saved, media_dir, media_path_prefix)

    if suffix in {".md", ".txt"}:
        return read_text(path) + note
    if suffix == ".csv":
        return read_csv_like(path, ",") + note
    if suffix == ".tsv":
        return read_csv_like(path, "\t") + note
    if suffix in {".xlsx", ".xlsm"}:
        return read_xlsx(path) + note
    if suffix == ".docx":
        return read_docx(path) + note
    if suffix == ".pdf":
        return read_pdf(path) + note
    raise RuntimeError(f"不支持的文件类型：{path.suffix}")


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract itinerary text from files.")
    parser.add_argument("inputs", nargs="+", help="行程文件路径")
    parser.add_argument("--output", "-o", help="输出文本文件；未提供则打印到 stdout")
    parser.add_argument("--media-dir", help="抽取文档内图片到指定目录；通常传新攻略目录/assets")
    parser.add_argument("--media-path-prefix", default=None, help="写入抽取文本中的图片路径前缀；默认使用 media-dir 目录名，攻略网页通常传 assets")
    args = parser.parse_args()
    media_dir = Path(args.media_dir).expanduser().resolve() if args.media_dir else None

    sections: list[str] = []
    for raw in args.inputs:
        path = Path(raw).expanduser().resolve()
        if not path.exists():
            raise SystemExit(f"文件不存在：{path}")
        try:
            body = extract(path, media_dir, args.media_path_prefix).strip()
        except Exception as exc:
            body = f"[抽取失败] {exc}"
        sections.append(f"\n\n===== {path.name} =====\n{body}\n")

    output = "\n".join(sections).strip() + "\n"
    if args.output:
        Path(args.output).write_text(output, encoding="utf-8")
    else:
        sys.stdout.write(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
